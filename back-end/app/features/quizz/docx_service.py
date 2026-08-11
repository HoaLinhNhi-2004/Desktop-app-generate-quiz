"""
DOCX Service - Extract text from DOCX files
"""

import os
import logging
import docx

logger = logging.getLogger(__name__)

# Word 97-2003 files are OLE2 compound documents; .docx is a ZIP of XML parts.
_OLE2_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"
_ZIP_MAGIC = b"PK\x03\x04"


def _reject_legacy_doc(path: str) -> None:
    """Fail early, and in the user's language, on a real Word 97-2003 file.

    `.doc` is accepted at every entry point but `python-docx` only opens OOXML,
    so a genuine legacy file surfaced as `PackageNotFoundError: Package not found
    at '...'` — which reads as a broken app, and Reprocess just repeats it. Still
    schools hand out `.doc` constantly, and a good share of those are really
    `.docx` renamed, which open fine — so sniff the magic bytes rather than
    refusing on the extension.
    """
    with open(path, "rb") as fh:
        header = fh.read(8)

    if header.startswith(_OLE2_MAGIC):
        raise ValueError(
            "File Word 97-2003 (.doc) không đọc được. Hãy mở bằng Word rồi chọn "
            "Save As → Word Document (.docx), sau đó tải lên lại."
        )
    if not header.startswith(_ZIP_MAGIC):
        raise ValueError(
            "File Word không hợp lệ hoặc đã hỏng. Hãy mở lại bằng Word và lưu "
            "thành .docx."
        )


def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract text from a DOCX file using python-docx.

    Args:
        docx_path: Path to the DOCX file
    
    Returns:
        Extracted text as a single string
    """
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    _reject_legacy_doc(docx_path)

    all_text = []
    try:
        doc = docx.Document(docx_path)
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                all_text.append(para.text.strip())
                
        # Extract from tables, one line per row with cells kept side by side.
        #
        # This used to append cell-by-cell, guarded by `cell.text not in all_text`
        # — a global dedupe over a list, so O(n²) *and* destructive: a 100-row
        # true/false bank collapsed to a single "Đúng" and a single "Sai" for the
        # whole document, which also drove `score_question_bank` to 0 so the
        # Import button never appeared for the very files it exists to serve.
        #
        # Only horizontally-merged repeats are collapsed: python-docx returns the
        # same cell object once per grid column it spans.
        for table in doc.tables:
            for row in table.rows:
                cells: list[str] = []
                seen_ids: set[int] = set()
                for cell in row.cells:
                    if id(cell._tc) in seen_ids:
                        continue
                    seen_ids.add(id(cell._tc))
                    cells.append(cell.text.strip())
                if any(cells):
                    all_text.append(" | ".join(cells))


    except Exception as e:
        logger.error(f"Error reading DOCX {docx_path}: {e}")
        raise

    combined = "\n\n".join(all_text)
    logger.info(f"Extracted {len(combined)} chars of text from {os.path.basename(docx_path)}")
    return combined
